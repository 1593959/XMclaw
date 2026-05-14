"""B-135 — ContentTools unit tests.

Pins:
  * list_tools advertises every tool when its enable flag is on
  * each tool degrades to ok=False with a clear "install X" hint
    when the optional dep is missing — never raises out of invoke()
  * happy paths for the deps we KNOW are bundled in the dev env
    (mss, python-docx, openpyxl, pyperclip, PIL)
  * pdf_read returns a structured "install pypdf" error since pypdf
    is not in the dev env (acts as the "missing-dep" pin)
"""
from __future__ import annotations

import base64
import importlib.util
import json
from pathlib import Path

import pytest

from xmclaw.core.ir import ToolCall
from xmclaw.providers.tool.content import ContentTools


def _call(name: str, args: dict | None = None) -> ToolCall:
    return ToolCall(
        name=name, args=args or {}, provenance="synthetic",
    )


# ── tool list ─────────────────────────────────────────────────────


def test_list_tools_default_includes_all() -> None:
    names = {s.name for s in ContentTools().list_tools()}
    assert names == {
        "screenshot", "image_read", "view_image", "view_video",
        "pdf_read", "docx_read", "xlsx_read",
        "clipboard_read", "clipboard_write",
    }


def test_list_tools_disable_screenshot() -> None:
    names = {s.name for s in ContentTools(enable_screenshot=False).list_tools()}
    assert "screenshot" not in names
    assert "image_read" in names  # always on


def test_list_tools_disable_clipboard() -> None:
    names = {s.name for s in ContentTools(enable_clipboard=False).list_tools()}
    assert "clipboard_read" not in names
    assert "clipboard_write" not in names


def test_list_tools_disable_documents() -> None:
    names = {s.name for s in ContentTools(enable_documents=False).list_tools()}
    assert "pdf_read" not in names
    assert "docx_read" not in names
    assert "xlsx_read" not in names


# ── unknown tool name ────────────────────────────────────────────


@pytest.mark.asyncio
async def test_unknown_tool_returns_structured_error() -> None:
    r = await ContentTools().invoke(_call("does_not_exist"))
    assert r.ok is False
    assert "unknown tool" in (r.error or "")


# ── pdf_read missing-dep pin ─────────────────────────────────────


@pytest.mark.asyncio
async def test_pdf_read_missing_pypdf_dep(tmp_path: Path) -> None:
    """pypdf isn't in the dev env (only pdfplumber/etc. would help).
    The tool must surface an actionable 'install pypdf' message
    rather than an ImportError."""
    if importlib.util.find_spec("pypdf"):
        pytest.skip("pypdf is installed; this test only fires without it")
    pdf = tmp_path / "x.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake")
    r = await ContentTools().invoke(_call("pdf_read", {"path": str(pdf)}))
    assert r.ok is False
    assert "pypdf" in (r.error or "")


# ── image_read happy + size cap ──────────────────────────────────


_TINY_PNG_HEX = (
    "89504E470D0A1A0A0000000D49484452000000010000000108060000001F15C4"
    "890000000A49444154789C636000000000020001E221BC330000000049454E44"
    "AE426082"
)


@pytest.mark.asyncio
async def test_image_read_defaults_to_no_base64(tmp_path: Path) -> None:
    """DEFAULT behaviour: NO base64 in the result (was the prompt-bloat
    bug). Path + metadata only, plus OCR text if an engine is wired."""
    img = tmp_path / "tiny.png"
    img.write_bytes(bytes.fromhex(_TINY_PNG_HEX))
    r = await ContentTools().invoke(_call("image_read", {"path": str(img)}))
    assert r.ok is True
    payload = json.loads(r.content)
    assert payload["mime"] == "image/png"
    assert payload["bytes"] == img.stat().st_size
    # The crucial assertion — no base64 by default.
    assert "base64" not in payload, (
        "image_read MUST NOT default-include base64 — stuffs ~MBs into "
        "the next prompt and the LLM can't read base64 from a tool result"
    )


@pytest.mark.asyncio
async def test_image_read_opt_in_base64(tmp_path: Path) -> None:
    """Explicit ``include_base64: true`` brings the bytes back."""
    img = tmp_path / "tiny.png"
    img.write_bytes(bytes.fromhex(_TINY_PNG_HEX))
    r = await ContentTools().invoke(_call(
        "image_read",
        {"path": str(img), "include_base64": True, "ocr": False},
    ))
    assert r.ok is True
    payload = json.loads(r.content)
    assert "base64" in payload
    decoded = base64.b64decode(payload["base64"])
    assert decoded == img.read_bytes()


@pytest.mark.asyncio
async def test_image_read_base64_skipped_when_file_too_big(
    tmp_path: Path,
) -> None:
    """Wave 25.7 regression: when include_base64=True but the file is
    above the 512 KB inline cap, we MUST skip the base64 payload (it
    would push the next LLM call past common API message-size limits).
    Vision attachment via metadata still happens — the LLM can still
    'see' the image through the multimodal pipeline."""
    img = tmp_path / "huge.png"
    # Build a 1 MB-ish file (above the 512 KB cap, below the 8 MB
    # hard cap). Starts with a valid PNG header so mime detection
    # still classifies it correctly.
    img.write_bytes(bytes.fromhex(_TINY_PNG_HEX) + b"\x00" * (700 * 1024))
    r = await ContentTools().invoke(_call(
        "image_read",
        {"path": str(img), "include_base64": True, "ocr": False},
    ))
    assert r.ok is True
    payload = json.loads(r.content)
    assert "base64" not in payload, (
        "image_read MUST NOT inline base64 for big files even when "
        "the caller asked for it — base64 in tool results blows up "
        "the next API call past message-size limits."
    )
    assert payload["base64_skipped"] is True
    assert "inline cap" in payload["base64_skip_reason"]
    # Vision attachment is preserved so the multimodal pipeline still
    # gives the LLM access to the image.
    assert payload["vision_attached"] is True
    assert r.metadata.get("attach_image") == str(img)


@pytest.mark.asyncio
async def test_image_read_ocr_skipped_when_flag_false(tmp_path: Path) -> None:
    img = tmp_path / "tiny.png"
    img.write_bytes(bytes.fromhex(_TINY_PNG_HEX))
    r = await ContentTools().invoke(_call(
        "image_read", {"path": str(img), "ocr": False},
    ))
    assert r.ok is True
    payload = json.loads(r.content)
    assert "ocr_text" not in payload
    assert "ocr_error" not in payload


@pytest.mark.asyncio
async def test_image_read_missing_file(tmp_path: Path) -> None:
    r = await ContentTools().invoke(_call(
        "image_read", {"path": str(tmp_path / "nope.png")},
    ))
    assert r.ok is False
    assert "not found" in (r.error or "")


@pytest.mark.asyncio
async def test_image_read_rejects_oversized(tmp_path: Path) -> None:
    big = tmp_path / "big.png"
    big.write_bytes(b"x" * (9 * 1024 * 1024))  # 9 MB
    r = await ContentTools().invoke(_call("image_read", {"path": str(big)}))
    assert r.ok is False
    assert "8 MB cap" in (r.error or "")


# ── docx_read happy ─────────────────────────────────────────────


@pytest.mark.asyncio
async def test_docx_read_extracts_paragraphs(tmp_path: Path) -> None:
    if not importlib.util.find_spec("docx"):
        pytest.skip("python-docx not installed")
    import docx as _docx  # type: ignore
    doc = _docx.Document()
    doc.add_paragraph("第一段 hello")
    doc.add_paragraph("第二段 world")
    p = tmp_path / "demo.docx"
    doc.save(str(p))

    r = await ContentTools().invoke(_call("docx_read", {"path": str(p)}))
    assert r.ok is True
    payload = json.loads(r.content)
    assert payload["paragraphs"] == 2
    assert "第一段 hello" in payload["text"]
    assert "第二段 world" in payload["text"]


# ── xlsx_read happy ─────────────────────────────────────────────


@pytest.mark.asyncio
async def test_xlsx_read_extracts_cells(tmp_path: Path) -> None:
    if not importlib.util.find_spec("openpyxl"):
        pytest.skip("openpyxl not installed")
    import openpyxl  # type: ignore
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["name", "value"])
    ws.append(["foo", 42])
    p = tmp_path / "demo.xlsx"
    wb.save(str(p))

    r = await ContentTools().invoke(_call("xlsx_read", {"path": str(p)}))
    assert r.ok is True
    payload = json.loads(r.content)
    assert "Sheet1" in payload["sheets"]
    assert "foo" in payload["text"]
    assert "42" in payload["text"]


# ── screenshot happy (mss is bundled) ────────────────────────────


@pytest.mark.asyncio
async def test_screenshot_writes_file(tmp_path: Path) -> None:
    if not importlib.util.find_spec("mss"):
        pytest.skip("mss not installed (CI without display)")
    out = tmp_path / "shot.png"
    r = await ContentTools().invoke(_call(
        "screenshot", {"path": str(out)},
    ))
    # Some headless environments (CI Linux without DISPLAY) raise from
    # mss. We accept both "ok=True with file" and "ok=False with
    # diagnostic" — what we DON'T accept is a Python exception leaking
    # out of invoke().
    if r.ok:
        assert out.is_file()
        payload = json.loads(r.content)
        assert payload["bytes"] == out.stat().st_size
        assert payload["width"] > 0
    else:
        assert r.error  # has a human-readable error string


# ── clipboard round-trip ────────────────────────────────────────


@pytest.mark.asyncio
async def test_clipboard_round_trip() -> None:
    if not importlib.util.find_spec("pyperclip"):
        pytest.skip("pyperclip not installed")
    tools = ContentTools()
    payload = "xmclaw clipboard test 123"
    write_r = await tools.invoke(_call("clipboard_write", {"text": payload}))
    if not write_r.ok:
        # Headless / no DISPLAY — pyperclip raises. That's a valid
        # not-installed-on-this-host case; just confirm the error is
        # diagnostic rather than a stack trace.
        pytest.skip(f"clipboard unavailable on this runner: {write_r.error}")

    read_r = await tools.invoke(_call("clipboard_read"))
    assert read_r.ok is True
    payload_back = json.loads(read_r.content)
    assert payload_back["text"] == payload


# ── Wave 26: view_image / view_video + MediaAttachment plumbing ──


@pytest.mark.asyncio
async def test_view_image_emits_canonical_attachments(
    tmp_path: Path,
) -> None:
    img = tmp_path / "tiny.png"
    img.write_bytes(bytes.fromhex(_TINY_PNG_HEX))
    r = await ContentTools().invoke(_call(
        "view_image", {"path": str(img)},
    ))
    assert r.ok is True
    payload = json.loads(r.content)
    assert payload["vision_attached"] is True
    # Legacy attach_image preserved for backwards compat.
    assert r.metadata["attach_image"] == str(img)
    # Canonical attachments list.
    atts = r.metadata["attachments"]
    assert len(atts) == 1
    assert atts[0]["kind"] == "image"
    assert atts[0]["path"] == str(img)
    assert atts[0]["mime"] == "image/png"


@pytest.mark.asyncio
async def test_view_image_rejects_unsupported_ext(
    tmp_path: Path,
) -> None:
    bad = tmp_path / "foo.txt"
    bad.write_text("not an image", encoding="utf-8")
    r = await ContentTools().invoke(_call("view_image", {"path": str(bad)}))
    assert r.ok is False
    assert "unsupported image extension" in r.error


@pytest.mark.asyncio
async def test_view_image_missing_file(tmp_path: Path) -> None:
    r = await ContentTools().invoke(_call(
        "view_image", {"path": str(tmp_path / "ghost.png")},
    ))
    assert r.ok is False
    assert "not found" in r.error


@pytest.mark.asyncio
async def test_view_video_basic_attachment(tmp_path: Path) -> None:
    """We don't need a real video — extension check + metadata
    structure is what we're locking in. cv2 duration probe just
    silently fails for the dummy bytes."""
    fake = tmp_path / "clip.mp4"
    # Tiny non-zero file; cv2 will fail to open but the tool returns
    # OK with duration_s=None.
    fake.write_bytes(b"\x00\x00\x00\x18ftypmp42" + b"\x00" * 64)
    r = await ContentTools().invoke(_call(
        "view_video", {"path": str(fake)},
    ))
    assert r.ok is True
    payload = json.loads(r.content)
    assert payload["mime"] == "video/mp4"
    assert payload["vision_attached"] is True
    atts = r.metadata["attachments"]
    assert len(atts) == 1
    assert atts[0]["kind"] == "video"
    assert atts[0]["mime"] == "video/mp4"
    # No legacy attach_image — video doesn't bridge into the old key.
    assert "attach_image" not in r.metadata


@pytest.mark.asyncio
async def test_view_video_rejects_unsupported_ext(
    tmp_path: Path,
) -> None:
    bad = tmp_path / "clip.flv"
    bad.write_bytes(b"FLV\x01" + b"\x00" * 16)
    r = await ContentTools().invoke(_call("view_video", {"path": str(bad)}))
    assert r.ok is False
    assert "unsupported video extension" in r.error


@pytest.mark.asyncio
async def test_screenshot_emits_canonical_attachments_too(
    tmp_path: Path, monkeypatch,
) -> None:
    """Regression: the screenshot tool's Wave 25.8 attach_image fix is
    now joined by the Wave 26 attachments list. Both must appear so
    legacy consumers + new normalize_attachments() consumers agree."""
    import mss
    if not hasattr(mss, "mss"):
        pytest.skip("mss not properly installed on runner")
    out = tmp_path / "shot.png"
    r = await ContentTools().invoke(_call(
        "screenshot", {"path": str(out), "monitor": 0},
    ))
    if not r.ok:
        # Headless CI / no DISPLAY — that's not the bug we're testing.
        pytest.skip(f"screenshot unavailable: {r.error}")
    # Legacy and canonical both present.
    assert r.metadata["attach_image"] == str(out)
    atts = r.metadata["attachments"]
    assert len(atts) == 1
    assert atts[0]["kind"] == "image"
    assert atts[0]["mime"] == "image/png"

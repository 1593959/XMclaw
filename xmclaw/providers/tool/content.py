"""ContentTools — desktop / clipboard / document-reader tools.

B-135. The user kept asking the agent to take a screenshot or read a
PDF and got "I don't have that tool". OpenClaw / Hermes / CoPaw all
ship these as built-ins; XMclaw's tool surface ended at file_read +
bash, which made the agent feel narrower than it actually is.

Tools shipped:

  * ``screenshot``       — capture the desktop (or a region) to a PNG
  * ``image_read``       — return image bytes as base64 so multimodal
                           LLMs can see what's there
  * ``pdf_read``         — extract text from a PDF (best-effort, falls
                           back to a clear "install pypdf" error)
  * ``docx_read``        — extract text from a .docx
  * ``xlsx_read``        — extract cells from a .xlsx (sheet by sheet)
  * ``clipboard_read``   — read the OS clipboard
  * ``clipboard_write``  — write text to the OS clipboard

Each tool gracefully reports a structured error when the underlying
optional dep is missing, so a fresh install with no extras gets a
diagnostic message instead of a crash. Heavy deps (pypdf for PDF,
python-docx for DOCX, openpyxl for XLSX) are imported inside the
handler — never at module import time — so the daemon doesn't fail
to boot on a partial install.

Layering note
-------------

Lives under ``xmclaw/providers/tool/`` alongside ``builtin.py`` per
``providers/tool/AGENTS.md`` §1. Composed into the agent's tool stack
in :file:`xmclaw/daemon/app.py` post-construction (same pattern as
agent_inter and the skill bridges).
"""
from __future__ import annotations

import base64
import json
import time
from pathlib import Path
from typing import Any

from xmclaw.core.ir import ToolCall, ToolResult, ToolSpec
from xmclaw.providers.tool.base import ToolProvider


# ── Specs ─────────────────────────────────────────────────────────


_SCREENSHOT_SPEC = ToolSpec(
    name="screenshot",
    description=(
        "Capture the user's desktop and save it to a PNG file. "
        "Returns the file path + image dimensions. The user can then "
        "view it directly, or you can call image_read to fetch the "
        "bytes for vision-capable LLMs.\n\n"
        "Use this when the user asks 'show me the screen', 'what's on "
        "my desktop', 'screenshot of the error', etc."
    ),
    parameters_schema={
        "type": "object",
        "properties": {
            "path": {
                "type": "string",
                "description": (
                    "Output PNG path. Defaults to "
                    "``~/Pictures/xmclaw_<unix_ts>.png`` on Windows or "
                    "``~/xmclaw_<unix_ts>.png`` elsewhere."
                ),
            },
            "monitor": {
                "type": "integer",
                "description": (
                    "Which monitor index to capture (1 = primary, 0 = "
                    "all monitors merged). Default 1."
                ),
            },
        },
    },
)


_IMAGE_READ_SPEC = ToolSpec(
    name="image_read",
    description=(
        "Read an image file (png/jpg/gif/webp) and return its OCR "
        "TEXT + metadata (path, mime, width, height, bytes). DEFAULT "
        "MODE is OCR-only — base64 is NOT returned by default because "
        "stuffing 3-5 MB of base64 into a tool result blows up the "
        "next LLM call (~1M tokens, 100s+ latency) AND the model can't "
        "actually 'see' base64 inside a tool result anyway — vision "
        "needs proper image content blocks. Set ``include_base64: "
        "true`` ONLY when you have downstream code that consumes it. "
        "Set ``ocr: false`` if you only need metadata. Capped at 8 MB."
    ),
    parameters_schema={
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "ocr": {
                "type": "boolean",
                "description": (
                    "Run OCR and include the extracted text. Default "
                    "true. Set false to skip OCR when you only need "
                    "metadata (path, size, dimensions)."
                ),
            },
            "include_base64": {
                "type": "boolean",
                "description": (
                    "Include the raw base64-encoded bytes in the "
                    "result. Default false. Almost always leave this "
                    "off — it explodes the prompt and the LLM cannot "
                    "interpret base64 from a tool result. Even when "
                    "set to true, the daemon enforces a 512 KB raw-"
                    "file cap; over that, base64 is silently skipped "
                    "(see base64_skipped flag in the result) so the "
                    "next LLM call doesn't get rejected for message "
                    "size. Vision attachment is independent and "
                    "always works via the multimodal pipeline."
                ),
            },
        },
        "required": ["path"],
    },
)


_PDF_READ_SPEC = ToolSpec(
    name="pdf_read",
    description=(
        "Extract text from a PDF, page by page. Returns text with "
        "'\\n--- page N ---\\n' separators. Defaults to the first 20 "
        "pages; use page_start/page_end for ranges, max_pages to "
        "widen the cap. Requires the optional ``pypdf`` dependency."
    ),
    parameters_schema={
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "page_start": {"type": "integer", "description": "1-indexed."},
            "page_end": {"type": "integer", "description": "Inclusive."},
            "max_pages": {"type": "integer", "description": "Cap; default 20."},
        },
        "required": ["path"],
    },
)


_DOCX_READ_SPEC = ToolSpec(
    name="docx_read",
    description=(
        "Extract paragraph text from a .docx file. Returns plain text "
        "with one paragraph per line; tables are flattened to "
        "tab-separated rows. Requires ``python-docx``."
    ),
    parameters_schema={
        "type": "object",
        "properties": {"path": {"type": "string"}},
        "required": ["path"],
    },
)


_XLSX_READ_SPEC = ToolSpec(
    name="xlsx_read",
    description=(
        "Read .xlsx workbooks. By default returns every sheet's cells "
        "as a CSV-like text block. Pass ``sheet`` to scope to one "
        "sheet, ``max_rows``/``max_cols`` to cap a large workbook. "
        "Requires ``openpyxl``."
    ),
    parameters_schema={
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "sheet": {"type": "string", "description": "Sheet name; omit for all."},
            "max_rows": {"type": "integer", "description": "Cap per sheet (default 200)."},
            "max_cols": {"type": "integer", "description": "Cap per sheet (default 30)."},
        },
        "required": ["path"],
    },
)


_CLIPBOARD_READ_SPEC = ToolSpec(
    name="clipboard_read",
    description=(
        "Read the OS clipboard's current text content. Useful when "
        "the user just copied a URL, error message, or snippet and "
        "wants the agent to act on it without pasting again."
    ),
    parameters_schema={"type": "object", "properties": {}},
)


_CLIPBOARD_WRITE_SPEC = ToolSpec(
    name="clipboard_write",
    description=(
        "Write text to the OS clipboard so the user can paste it "
        "elsewhere. Use after producing a snippet, command, or URL "
        "the user will likely want to copy."
    ),
    parameters_schema={
        "type": "object",
        "properties": {
            "text": {"type": "string"},
        },
        "required": ["text"],
    },
)


# ── Provider ──────────────────────────────────────────────────────


class ContentTools(ToolProvider):
    """Desktop / clipboard / document tools.

    No constructor knobs — every tool degrades gracefully when its
    optional dep is missing or the OS doesn't support it (e.g.
    headless Linux + no DISPLAY for screenshot). The ``enable_*``
    flags exist so a security-conscious deployment can hide tools
    they don't want the LLM reaching for.
    """

    def __init__(
        self,
        *,
        enable_screenshot: bool = True,
        enable_clipboard: bool = True,
        enable_documents: bool = True,
    ) -> None:
        self._enable_screenshot = enable_screenshot
        self._enable_clipboard = enable_clipboard
        self._enable_documents = enable_documents

    def list_tools(self) -> list[ToolSpec]:
        out: list[ToolSpec] = []
        if self._enable_screenshot:
            out.append(_SCREENSHOT_SPEC)
        # image_read is always on — it's just file reading + base64
        out.append(_IMAGE_READ_SPEC)
        if self._enable_documents:
            out.extend([_PDF_READ_SPEC, _DOCX_READ_SPEC, _XLSX_READ_SPEC])
        if self._enable_clipboard:
            out.extend([_CLIPBOARD_READ_SPEC, _CLIPBOARD_WRITE_SPEC])
        return out

    async def invoke(self, call: ToolCall) -> ToolResult:
        t0 = time.perf_counter()
        name = call.name
        try:
            if name == "screenshot":
                return await self._screenshot(call, t0)
            if name == "image_read":
                return await self._image_read(call, t0)
            if name == "pdf_read":
                return await self._pdf_read(call, t0)
            if name == "docx_read":
                return await self._docx_read(call, t0)
            if name == "xlsx_read":
                return await self._xlsx_read(call, t0)
            if name == "clipboard_read":
                return await self._clipboard_read(call, t0)
            if name == "clipboard_write":
                return await self._clipboard_write(call, t0)
        except Exception as exc:  # noqa: BLE001 — every tool surfaces as ok=False
            return _fail(call, t0, f"{type(exc).__name__}: {exc}")
        return _fail(call, t0, f"unknown tool: {name!r}")

    # ── individual handlers ──────────────────────────────────────

    async def _screenshot(self, call: ToolCall, t0: float) -> ToolResult:
        try:
            import mss
        except ImportError:
            return _fail(call, t0, (
                "screenshot tool needs the ``mss`` package. "
                "Install with: pip install mss"
            ))
        args = call.args or {}
        # Default output path. ~/Pictures on Windows is the standard
        # screenshot dump; fall back to ~ on other OSes.
        ts = int(time.time())
        out_path_arg = args.get("path")
        if out_path_arg:
            out = Path(str(out_path_arg)).expanduser()
        else:
            home = Path.home()
            pics = home / "Pictures"
            base = pics if pics.is_dir() else home
            out = base / f"xmclaw_{ts}.png"
        out.parent.mkdir(parents=True, exist_ok=True)

        monitor_idx = int(args.get("monitor", 1))
        with mss.mss() as sct:
            mons = sct.monitors  # [0] = all, [1+] = individual
            if monitor_idx < 0 or monitor_idx >= len(mons):
                return _fail(call, t0, (
                    f"invalid monitor index {monitor_idx}; "
                    f"available: 0..{len(mons) - 1}"
                ))
            grab = sct.grab(mons[monitor_idx])
            mss.tools.to_png(grab.rgb, grab.size, output=str(out))
            w, h = grab.size

        return _ok(call, t0, json.dumps({
            "path": str(out),
            "width": w,
            "height": h,
            "monitor": monitor_idx,
            "bytes": out.stat().st_size,
        }, ensure_ascii=False))

    async def _image_read(self, call: ToolCall, t0: float) -> ToolResult:
        path = _path_arg(call)
        if isinstance(path, ToolResult):
            return path
        if not path.is_file():
            return _fail(call, t0, f"image not found: {path}")
        size = path.stat().st_size
        if size > 8 * 1024 * 1024:
            return _fail(call, t0, (
                f"image is {size / 1024 / 1024:.1f} MB — over the 8 MB cap. "
                f"Resize first or pass a smaller file."
            ))
        args = call.args or {}
        want_ocr = bool(args.get("ocr", True))
        want_b64 = bool(args.get("include_base64", False))

        ext = path.suffix.lower().lstrip(".")
        mime_map = {
            "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
            "gif": "image/gif", "webp": "image/webp", "bmp": "image/bmp",
        }
        mime = mime_map.get(ext, "application/octet-stream")

        payload: dict[str, Any] = {
            "path": str(path),
            "mime": mime,
            "bytes": size,
        }

        # Image dimensions — cheap (Pillow only reads the header).
        try:
            from PIL import Image  # type: ignore
            with Image.open(str(path)) as im:
                payload["width"], payload["height"] = im.size
        except Exception:  # noqa: BLE001 — Pillow optional; metadata is a nice-to-have
            pass

        if want_ocr:
            ocr_text, ocr_error = _run_ocr_on_path(str(path))
            if ocr_text is not None:
                payload["ocr_text"] = ocr_text
                payload["ocr_block_count"] = ocr_text.count("\n") + 1 if ocr_text else 0
            else:
                payload["ocr_error"] = ocr_error

        # Wave 25.7: even with include_base64=true, refuse to inline
        # giants. ~512 KB raw inflates to ~700 KB base64 — that fits
        # under common 2 MB API message caps even when stacked with
        # tool history. Bigger than that → drop the base64, keep the
        # metadata + attach_image so vision pipeline still works.
        _BASE64_INLINE_CAP = 512 * 1024
        if want_b64:
            if size > _BASE64_INLINE_CAP:
                payload["base64_skipped"] = True
                payload["base64_skip_reason"] = (
                    f"file is {size / 1024:.0f} KB, over the "
                    f"{_BASE64_INLINE_CAP // 1024} KB inline cap. "
                    "Vision is still attached via the multimodal "
                    "pipeline (see vision_attached). The LLM cannot "
                    "read base64 from a tool result anyway — request "
                    "an image-aware turn instead."
                )
            else:
                data = path.read_bytes()
                payload["base64"] = base64.b64encode(data).decode("ascii")

        payload["vision_attached"] = True
        return ToolResult(
            call_id=call.id, ok=True,
            content=json.dumps(payload, ensure_ascii=False),
            latency_ms=(time.perf_counter() - t0) * 1000.0,
            metadata={"attach_image": str(path)},
        )

    async def _pdf_read(self, call: ToolCall, t0: float) -> ToolResult:
        try:
            import pypdf  # type: ignore
        except ImportError:
            return _fail(call, t0, (
                "pdf_read tool needs the ``pypdf`` package. "
                "Install with: pip install pypdf"
            ))
        path = _path_arg(call)
        if isinstance(path, ToolResult):
            return path
        if not path.is_file():
            return _fail(call, t0, f"pdf not found: {path}")
        args = call.args or {}
        max_pages = int(args.get("max_pages", 20))
        page_start = int(args.get("page_start", 1))
        page_end_arg = args.get("page_end")

        reader = pypdf.PdfReader(str(path))
        total = len(reader.pages)
        page_end = (
            int(page_end_arg) if page_end_arg is not None
            else min(total, page_start + max_pages - 1)
        )
        page_start = max(1, page_start)
        page_end = min(total, page_end)
        out_chunks: list[str] = []
        for i in range(page_start - 1, page_end):
            try:
                txt = reader.pages[i].extract_text() or ""
            except Exception as exc:  # noqa: BLE001
                txt = f"<extract failed: {exc}>"
            out_chunks.append(f"\n--- page {i + 1} ---\n{txt.strip()}")
        body = "\n".join(out_chunks).strip()
        if not body:
            body = "(pdf has no extractable text — likely scanned image)"
        return _ok(call, t0, json.dumps({
            "path": str(path),
            "total_pages": total,
            "pages_read": [page_start, page_end],
            "text": body[:32000],  # cap
        }, ensure_ascii=False))

    async def _docx_read(self, call: ToolCall, t0: float) -> ToolResult:
        try:
            import docx
        except ImportError:
            return _fail(call, t0, (
                "docx_read needs ``python-docx``. "
                "Install with: pip install python-docx"
            ))
        path = _path_arg(call)
        if isinstance(path, ToolResult):
            return path
        if not path.is_file():
            return _fail(call, t0, f"docx not found: {path}")
        d = docx.Document(str(path))
        out_lines: list[str] = []
        for p in d.paragraphs:
            txt = p.text.strip()
            if txt:
                out_lines.append(txt)
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    out_lines.append("\t".join(cells))
        return _ok(call, t0, json.dumps({
            "path": str(path),
            "paragraphs": len(d.paragraphs),
            "tables": len(d.tables),
            "text": "\n".join(out_lines)[:32000],
        }, ensure_ascii=False))

    async def _xlsx_read(self, call: ToolCall, t0: float) -> ToolResult:
        try:
            import openpyxl  # type: ignore
        except ImportError:
            return _fail(call, t0, (
                "xlsx_read needs ``openpyxl``. "
                "Install with: pip install openpyxl"
            ))
        path = _path_arg(call)
        if isinstance(path, ToolResult):
            return path
        if not path.is_file():
            return _fail(call, t0, f"xlsx not found: {path}")
        args = call.args or {}
        max_rows = int(args.get("max_rows", 200))
        max_cols = int(args.get("max_cols", 30))
        sheet_arg = args.get("sheet")

        # data_only=True so formulas resolve to their cached values
        # rather than returning the formula string.
        wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        try:
            sheets = (
                [sheet_arg] if sheet_arg
                else wb.sheetnames
            )
            out_sections: list[str] = []
            for sname in sheets:
                if sname not in wb.sheetnames:
                    out_sections.append(f"\n=== {sname} ===\n(sheet not found)")
                    continue
                ws = wb[sname]
                out_sections.append(f"\n=== {sname} ===")
                for r_idx, row in enumerate(ws.iter_rows(values_only=True)):
                    if r_idx >= max_rows:
                        out_sections.append(f"... [truncated at {max_rows} rows]")
                        break
                    cells = [
                        ("" if v is None else str(v))
                        for v in row[:max_cols]
                    ]
                    out_sections.append("\t".join(cells))
        finally:
            wb.close()
        return _ok(call, t0, json.dumps({
            "path": str(path),
            "sheets": wb.sheetnames,
            "text": "\n".join(out_sections)[:32000],
        }, ensure_ascii=False))

    async def _clipboard_read(self, call: ToolCall, t0: float) -> ToolResult:
        try:
            import pyperclip  # type: ignore
        except ImportError:
            return _fail(call, t0, (
                "clipboard_read needs ``pyperclip``. "
                "Install with: pip install pyperclip"
            ))
        try:
            text = pyperclip.paste()
        except Exception as exc:  # pyperclip raises on no display
            return _fail(call, t0, f"clipboard unavailable: {exc}")
        if not isinstance(text, str):
            text = str(text or "")
        # Guard: a 50 MB clipboard would blow context. Cap at 32 KB.
        if len(text) > 32_000:
            text = text[:32_000] + "\n[truncated, clipboard had more]"
        return _ok(call, t0, json.dumps({
            "text": text,
            "length": len(text),
        }, ensure_ascii=False))

    async def _clipboard_write(self, call: ToolCall, t0: float) -> ToolResult:
        try:
            import pyperclip
        except ImportError:
            return _fail(call, t0, (
                "clipboard_write needs ``pyperclip``. "
                "Install with: pip install pyperclip"
            ))
        text = (call.args or {}).get("text", "")
        if not isinstance(text, str):
            text = str(text)
        try:
            pyperclip.copy(text)
        except Exception as exc:
            return _fail(call, t0, f"clipboard write failed: {exc}")
        return _ok(call, t0, json.dumps({
            "ok": True,
            "length": len(text),
        }, ensure_ascii=False))


# ── helpers ───────────────────────────────────────────────────────


def _path_arg(call: ToolCall) -> Path | ToolResult:
    raw = (call.args or {}).get("path")
    if not isinstance(raw, str) or not raw.strip():
        return _fail(call, time.perf_counter(), "path argument required")
    return Path(raw).expanduser()


def _ok(call: ToolCall, t0: float, content: Any) -> ToolResult:
    return ToolResult(
        call_id=call.id, ok=True, content=content,
        latency_ms=(time.perf_counter() - t0) * 1000.0,
    )


def _fail(call: ToolCall, t0: float, err: str) -> ToolResult:
    return ToolResult(
        call_id=call.id, ok=False, content=None, error=err,
        latency_ms=(time.perf_counter() - t0) * 1000.0,
    )


def _run_ocr_on_path(image_path: str) -> tuple[str | None, str]:
    """Best-effort OCR over a saved image file.

    Returns ``(text, "")`` on success or ``(None, error_message)`` on
    any failure. Tries engines in order: rapidocr → paddleocr →
    pytesseract. All three accept either file paths or PIL images;
    we feed paths to keep the surface tiny.

    The returned ``text`` is one OCR block per line. Empty string is
    valid (image had no text). ``None`` means OCR genuinely failed —
    no engine installed, or all engines errored.
    """
    # rapidocr — best Chinese support per MB.
    try:
        from rapidocr_onnxruntime import RapidOCR  # type: ignore
        engine = RapidOCR()
        result, _ = engine(image_path)
        if result is None:
            return "", ""
        lines: list[str] = []
        for row in result:
            if not row or len(row) < 2:
                continue
            text = str(row[1] or "").strip()
            if text:
                lines.append(text)
        return "\n".join(lines), ""
    except ImportError:
        pass
    except Exception as exc:  # noqa: BLE001
        return None, f"rapidocr failed: {type(exc).__name__}: {exc}"

    # paddleocr.
    try:
        from paddleocr import PaddleOCR  # type: ignore
        ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
        result = ocr.ocr(image_path, cls=True)
        if not result:
            return "", ""
        lines = []
        for page in result:
            if not page:
                continue
            for row in page:
                if not row or len(row) < 2:
                    continue
                text_pair = row[1]
                text = str(text_pair[0] if isinstance(text_pair, (list, tuple)) else text_pair).strip()
                if text:
                    lines.append(text)
        return "\n".join(lines), ""
    except ImportError:
        pass
    except Exception as exc:  # noqa: BLE001
        return None, f"paddleocr failed: {type(exc).__name__}: {exc}"

    # pytesseract — needs the Tesseract binary on PATH.
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
        text = pytesseract.image_to_string(Image.open(image_path), lang="chi_sim+eng")
        return (text or "").strip(), ""
    except ImportError:
        pass
    except Exception as exc:  # noqa: BLE001
        return None, f"pytesseract failed: {type(exc).__name__}: {exc}"

    return None, (
        "no OCR engine available — install one of: "
        "rapidocr-onnxruntime / paddleocr / pytesseract"
    )
